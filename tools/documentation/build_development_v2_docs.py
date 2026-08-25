from __future__ import annotations

import csv
import hashlib
import re
import shutil
import textwrap
from datetime import date
from pathlib import Path

import arabic_reshaper
from bidi.algorithm import get_display
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "deliverables" / "Project-Desk-Development-v2"
SOURCE_OUT = OUT / "source"
DIAGRAM_OUT = OUT / "diagrams"
MATRIX_OUT = OUT / "matrices"

BRAND_CYAN = "16C8CE"
BRAND_NAVY = "406286"
BRAND_DARK = "243442"
BRAND_LIGHT = "F3F7F8"
BRAND_MUTED = "6A7782"
BRAND_ORANGE = "E59E37"
WHITE = "FFFFFF"

FONT_AR = "Arial"
FONT_CODE = "Consolas"
TODAY_AR = "23 أغسطس 2026"
VERSION = "2.0-development"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig").strip()


def compose_sources() -> dict[str, Path]:
    SOURCE_OUT.mkdir(parents=True, exist_ok=True)
    old_srs = read_text(ROOT / "docs" / "specification" / "PROJECT_DESK_SRS_IEEE830_AR.md")
    srs_add = read_text(ROOT / "docs" / "development-v2" / "SRS_V2_ADDENDUM_AR.md")
    srs = (
        "# Project Desk — كراسة مواصفات المتطلبات الشاملة (IEEE 830)\n\n"
        "> النسخة: 2.0 Development — وثيقة شاملة تحفظ خط أساس النسخة الأولى وتدمج إضافات نسخة التطوير.\n\n"
        "## طريقة قراءة هذه الكراسة\n\n"
        "يتكوّن هذا المستند من خط الأساس الكامل المعتمد سابقًا، ثم ملحق تطوير إلزامي يحدّثه. "
        "عند التعارض، تكون متطلبات ملحق نسخة التطوير هي السائدة داخل فرع `develop` فقط.\n\n"
        "# الجزء الأول: خط الأساس الكامل للنسخة الأولى\n\n"
        + old_srs
        + "\n\n# الجزء الثاني: متطلبات نسخة التطوير 2.0\n\n"
        + srs_add
    )

    tech_files = [
        "DOCUMENTATION_INDEX_AR.md",
        "SYSTEM_OVERVIEW_AR.md",
        "ARCHITECTURE_AR.md",
        "DATA_MODEL_AR.md",
        "API_AND_ROUTES_AR.md",
        "SECURITY_AND_PERMISSIONS_AR.md",
        "DEPLOYMENT_AR.md",
        "OPERATIONS_AND_RECOVERY_AR.md",
        "TESTING_AND_QA_AR.md",
        "DEVELOPER_GUIDE_AR.md",
    ]
    tech_parts = []
    for file_name in tech_files:
        title = file_name.removesuffix("_AR.md").replace("_", " ")
        tech_parts.append(f"# مرجع النسخة الأولى: {title}\n\n{read_text(ROOT / 'docs' / 'technical' / file_name)}")
    technical = (
        "# Project Desk — التوثيق التقني الشامل لنسخة التطوير\n\n"
        "> يشمل بنية النسخة الأولى كاملة ثم تكاملات المراحل، المشروع القائم، شجرة المتطلبات، والتحليل المحلي.\n\n"
        + "\n\n".join(tech_parts)
        + "\n\n# دليل التحليل المحلي للمتطلبات\n\n"
        + read_text(ROOT / "docs" / "LOCAL_REQUIREMENTS_ANALYSIS_AR.md")
        + "\n\n# ملحق البنية والتنفيذ لنسخة التطوير 2.0\n\n"
        + read_text(ROOT / "docs" / "development-v2" / "TECHNICAL_V2_ADDENDUM_AR.md")
    )

    manual = (
        "# Project Desk — دليل المستخدم الشامل لنسخة التطوير\n\n"
        "> يضم دليل النسخة الأولى كاملًا، ثم خطوات تشغيل الوظائف الجديدة في نسخة التطوير.\n\n"
        "# الجزء الأول: دليل النسخة الأولى\n\n"
        + read_text(ROOT / "docs" / "USER_MANUAL_AR.md")
        + "\n\n# الجزء الثاني: دليل وظائف نسخة التطوير 2.0\n\n"
        + read_text(ROOT / "docs" / "development-v2" / "USER_MANUAL_V2_ADDENDUM_AR.md")
    )

    research_root = ROOT.parent / "research" / "project-desk-global-comparison-2026-08"
    competitive = (
        "# Project Desk — المقارنة التنافسية الشاملة (خط أساس + نسخة التطوير)\n\n"
        "> تحفظ هذه الوثيقة نتائج الدراسة المرجعية المؤرخة في أغسطس 2026، وتحدّث تقييم Project Desk فقط "
        "وفق الوظائف المنفذة في فرع التطوير.\n\n"
        "# الجزء الأول: الدراسة المرجعية الكاملة\n\n"
        + read_text(research_root / "report-ar.md")
        + "\n\n# الجزء الثاني: أثر نسخة التطوير على المقارنة\n\n"
        + read_text(ROOT / "docs" / "development-v2" / "COMPETITIVE_V2_ADDENDUM_AR.md")
    )

    workflows = read_text(ROOT / "docs" / "development-v2" / "WORKFLOWS_V2_AR.md")
    docs = {
        "Project-Desk-Development-v2-SRS-IEEE830-AR.md": srs,
        "Project-Desk-Development-v2-Technical-Documentation-AR.md": technical,
        "Project-Desk-Development-v2-User-Manual-AR.md": manual,
        "Project-Desk-Development-v2-Workflows-AR.md": workflows,
        "Project-Desk-Development-v2-Competitive-Analysis-AR.md": competitive,
    }
    paths: dict[str, Path] = {}
    for name, content in docs.items():
        path = SOURCE_OUT / name
        path.write_text(content.rstrip() + "\n", encoding="utf-8")
        paths[name] = path

    for file_name in ["IMPLEMENTATION_VERIFICATION_AR.md", "DOCUMENTATION_INDEX_AR.md"]:
        shutil.copy2(ROOT / "docs" / "development-v2" / file_name, OUT / file_name)
    shutil.copy2(ROOT / "docs" / "development-v2" / "DOCUMENTATION_INDEX_AR.md", OUT / "README-FIRST-AR.md")

    MATRIX_OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ROOT / "docs" / "specification" / "REQUIREMENTS_MATRIX.csv", MATRIX_OUT / "requirements-matrix-v1-baseline.csv")
    for file_name in ["feature-gap-matrix.csv", "project-desk-advantages.csv", "project-desk-scorecard.csv", "system-landscape.csv"]:
        shutil.copy2(research_root / file_name, MATRIX_OUT / file_name)
    return paths


def ar_display(text: str) -> str:
    return get_display(arabic_reshaper.reshape(text))


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/tahomabd.ttf" if bold else "C:/Windows/Fonts/tahoma.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: ImageFont.FreeTypeFont, fill: str) -> None:
    x, y = xy
    shown = ar_display(text)
    box = draw.textbbox((0, 0), shown, font=font)
    draw.text((x - (box[2] - box[0]) / 2, y - (box[3] - box[1]) / 2), shown, font=font, fill=fill)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = BRAND_NAVY) -> None:
    draw.line([start, end], fill=f"#{fill}", width=6)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x2 - 22 * ux + 11 * px, y2 - 22 * uy + 11 * py)
    p2 = (x2 - 22 * ux - 11 * px, y2 - 22 * uy - 11 * py)
    draw.polygon([end, p1, p2], fill=f"#{fill}")


def make_flow_diagram(name: str, title: str, nodes: list[str], accent_indexes: set[int] | None = None) -> Path:
    accent_indexes = accent_indexes or set()
    width, height = 1800, 760
    img = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title_font = get_font(52, True)
    node_font = get_font(29, True)
    sub_font = get_font(24)
    draw.rounded_rectangle((35, 30, width - 35, height - 30), radius=35, outline=f"#{BRAND_CYAN}", width=5, fill=f"#{BRAND_LIGHT}")
    draw_centered(draw, (width // 2, 92), title, title_font, f"#{BRAND_DARK}")
    count = len(nodes)
    gap = 42
    box_w = int((width - 160 - gap * (count - 1)) / count)
    box_h = 250
    y = 245
    for i, label in enumerate(nodes):
        x = 80 + i * (box_w + gap)
        fill = BRAND_CYAN if i in accent_indexes else WHITE
        text_fill = WHITE if i in accent_indexes else BRAND_DARK
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=24, fill=f"#{fill}", outline=f"#{BRAND_NAVY}", width=4)
        wrapped = textwrap.wrap(label, width=max(9, 18 - count))
        start_y = y + box_h // 2 - (len(wrapped) * 38) // 2
        for j, line in enumerate(wrapped):
            draw_centered(draw, (x + box_w // 2, start_y + j * 44), line, node_font if len(wrapped) <= 2 else sub_font, f"#{text_fill}")
        if i < count - 1:
            arrow(draw, (x + box_w + 5, y + box_h // 2), (x + box_w + gap - 5, y + box_h // 2))
    path = DIAGRAM_OUT / name
    img.save(path, dpi=(180, 180))
    return path


def make_taxonomy_diagram() -> Path:
    width, height = 1800, 980
    img = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title_font = get_font(52, True)
    box_font = get_font(30, True)
    small_font = get_font(24)
    draw.rounded_rectangle((35, 30, width - 35, height - 30), radius=35, outline=f"#{BRAND_CYAN}", width=5, fill=f"#{BRAND_LIGHT}")
    draw_centered(draw, (width // 2, 90), "شجرة المتطلبات والعلاقات", title_font, f"#{BRAND_DARK}")
    project = (650, 170, 1150, 290)
    category1 = (260, 390, 700, 510)
    category2 = (1100, 390, 1540, 510)
    group1 = (170, 650, 610, 770)
    group2 = (690, 650, 1110, 770)
    req = (1190, 650, 1630, 770)
    boxes = [(project, "المشروع", BRAND_NAVY, WHITE), (category1, "فئة الأعمال", BRAND_CYAN, WHITE), (category2, "الفئة التقنية", BRAND_CYAN, WHITE), (group1, "مجموعة العقود", WHITE, BRAND_DARK), (group2, "مجموعة الفواتير", WHITE, BRAND_DARK), (req, "متطلبات التكامل", WHITE, BRAND_DARK)]
    for box, label, fill, text_fill in boxes:
        draw.rounded_rectangle(box, radius=22, fill=f"#{fill}", outline=f"#{BRAND_NAVY}", width=4)
        draw_centered(draw, ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2), label, box_font, f"#{text_fill}")
    arrow(draw, (760, 290), (530, 390))
    arrow(draw, (1040, 290), (1270, 390))
    arrow(draw, (480, 510), (390, 650))
    arrow(draw, (600, 510), (900, 650))
    arrow(draw, (1320, 510), (1410, 650))
    draw_centered(draw, (900, 870), "العلاقات: يعتمد على • يكمّل • يفصّل • يتعارض • مكرر • يستبدل • مرتبط بـ", small_font, f"#{BRAND_MUTED}")
    path = DIAGRAM_OUT / "05-requirement-taxonomy.png"
    img.save(path, dpi=(180, 180))
    return path


def make_architecture_diagram() -> Path:
    width, height = 1800, 1040
    img = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title_font = get_font(52, True)
    box_font = get_font(28, True)
    small_font = get_font(22)
    draw.rounded_rectangle((35, 30, width - 35, height - 30), radius=35, outline=f"#{BRAND_CYAN}", width=5, fill=f"#{BRAND_LIGHT}")
    draw_centered(draw, (width // 2, 85), "البنية المنطقية لنسخة التطوير", title_font, f"#{BRAND_DARK}")
    layers = [
        (150, 170, 1650, 310, "واجهة React + Inertia — عربي/إنجليزي — RTL/LTR", BRAND_NAVY, WHITE),
        (150, 390, 1650, 570, "Laravel: المشاريع • المراحل • المتطلبات • المرشحون • الصلاحيات", WHITE, BRAND_DARK),
        (150, 660, 770, 850, "قاعدة البيانات والطابور\nSQLite / Database Queue", WHITE, BRAND_DARK),
        (1030, 660, 1650, 850, "حد الثقة المحلي\nOllama + Qwen3 + OCR", BRAND_CYAN, WHITE),
    ]
    for x1, y1, x2, y2, label, fill, text_fill in layers:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=f"#{fill}", outline=f"#{BRAND_NAVY}", width=4)
        for j, line in enumerate(label.split("\n")):
            draw_centered(draw, ((x1+x2)//2, (y1+y2)//2 + (j - (len(label.split(chr(10)))-1)/2)*42), line, box_font, f"#{text_fill}")
    arrow(draw, (900, 310), (900, 390))
    arrow(draw, (650, 570), (490, 660))
    arrow(draw, (1150, 570), (1310, 660))
    draw_centered(draw, (900, 940), "لا يُنشئ النموذج متطلبات معتمدة مباشرة؛ جميع المخرجات تمر بمراجعة بشرية", small_font, f"#{BRAND_MUTED}")
    path = DIAGRAM_OUT / "01-development-architecture.png"
    img.save(path, dpi=(180, 180))
    return path


def create_diagrams() -> list[Path]:
    DIAGRAM_OUT.mkdir(parents=True, exist_ok=True)
    diagrams = [make_architecture_diagram()]
    diagrams.append(make_flow_diagram("02-existing-project-wizard.png", "إدخال مشروع قائم دون فقد تاريخه", ["البيانات والفريق", "تاريخ البداية والانتقال", "المراحل والمعالم", "المهام والمخاطر", "مراجعة واعتماد"], {4}))
    diagrams.append(make_flow_diagram("03-phase-milestone-governance.png", "المراحل الموزونة ومعالم التسليم", ["خطة بأوزان 100%", "تنفيذ المهام", "معلم إلزامي", "اعتماد المرحلة", "تقدم المشروع"], {2, 4}))
    diagrams.append(make_flow_diagram("04-local-ai-pipeline.png", "خط التحليل المحلي الآمن", ["رفع وفحص", "استخراج/OCR", "تقسيم المقاطع", "Ollama محلي", "مرشحون", "مراجعة بشرية"], {3, 5}))
    diagrams.append(make_taxonomy_diagram())
    diagrams.append(make_flow_diagram("06-candidate-review.png", "دورة مراجعة المرشحين", ["نتائج منظمة", "تحقق من المصدر", "تعديل/دمج/رفض", "اعتماد جماعي", "متطلبات وعلاقات"], {1, 3}))
    diagrams.append(make_flow_diagram("07-version-comparison.png", "مقارنة إصدارات الكراسة", ["إصدار معتمد", "إصدار جديد", "جديد/معدل/محذوف", "تحليل الأثر", "قرار بشري"], {2, 4}))
    return diagrams


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_rtl_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_rtl(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.32)
        section.different_first_page_header_footer = True


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_AR
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BRAND_DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    normal.element.rPr.rFonts.set(qn("w:ascii"), FONT_AR)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_AR)
    normal.element.rPr.rFonts.set(qn("w:cs"), FONT_AR)
    for idx, size, color in [(1, 18, BRAND_NAVY), (2, 14, BRAND_CYAN), (3, 12, BRAND_NAVY), (4, 11, BRAND_DARK)]:
        style = styles[f"Heading {idx}"]
        style.font.name = FONT_AR
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if idx <= 2 else 9)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True
        style.element.rPr.rFonts.set(qn("w:ascii"), FONT_AR)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_AR)
        style.element.rPr.rFonts.set(qn("w:cs"), FONT_AR)
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_AR
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.element.rPr.rFonts.set(qn("w:cs"), FONT_AR)


def setup_headers_footers(doc: Document, short_title: str) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.clear()
        set_rtl_paragraph(p)
        run = p.add_run(f"Project Desk | {short_title} | نسخة التطوير 2.0")
        run.font.name = FONT_AR
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
        set_run_rtl(run)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("CloudTech — وثيقة تطوير داخلية   •   ")
        r1.font.name = FONT_AR
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor.from_string(BRAND_MUTED)
        add_field(fp, "PAGE")


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CLOUDTECH")
    run.font.name = FONT_AR
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BRAND_CYAN)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(35)
    r2 = p2.add_run("PROJECT DESK")
    r2.font.name = FONT_AR
    r2.font.size = Pt(30)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(18)
    r3 = p3.add_run(title)
    r3.font.name = FONT_AR
    r3.font.size = Pt(23)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor.from_string(BRAND_DARK)
    set_run_rtl(r3)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(subtitle)
    r4.font.name = FONT_AR
    r4.font.size = Pt(12)
    r4.font.color.rgb = RGBColor.from_string(BRAND_MUTED)
    set_run_rtl(r4)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = [("الحالة", "مرجع نسخة التطوير"), ("الإصدار", VERSION), ("التاريخ", TODAY_AR), ("النطاق", "فرع develop — مستقل عن نسخة الاستخدام الحالية")]
    for row, (label, value) in zip(table.rows, labels):
        set_cell_width(row.cells[0], 1800)
        set_cell_width(row.cells[1], 5200)
        set_cell_shading(row.cells[0], BRAND_NAVY)
        for idx, text in enumerate([label, value]):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pcell = cell.paragraphs[0]
            set_rtl_paragraph(pcell)
            rr = pcell.add_run(text)
            rr.font.name = FONT_AR
            rr.font.size = Pt(10)
            rr.font.bold = idx == 0
            rr.font.color.rgb = RGBColor.from_string(WHITE if idx == 0 else BRAND_DARK)
            set_run_rtl(rr)
    set_repeat_table_header(table.rows[0])
    notice = doc.add_paragraph()
    notice.paragraph_format.space_before = Pt(22)
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = notice.add_run("هذه الحزمة توثّق نسخة التطوير فقط ولا تغيّر مستندات خط الإنتاج أو قاعدة بيانات الاستخدام الحالية.")
    rn.font.name = FONT_AR
    rn.font.size = Pt(9.5)
    rn.font.italic = True
    rn.font.color.rgb = RGBColor.from_string(BRAND_ORANGE)
    set_run_rtl(rn)
    doc.add_page_break()


INLINE_RE = re.compile(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))")


def add_inline(paragraph, text: str, *, rtl: bool = True) -> None:
    text = text.replace("<br>", "\n").replace("<br/>", "\n")
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.name = FONT_AR
            if rtl:
                set_run_rtl(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = FONT_AR
            if rtl:
                set_run_rtl(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = FONT_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
        else:
            m = re.match(r"\[(.*?)\]\((.*?)\)", token)
            label, url = m.groups() if m else (token, "")
            run = paragraph.add_run(f"{label} ({url})")
            run.font.name = FONT_AR
            run.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
            run.underline = True
            if rtl:
                set_run_rtl(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = FONT_AR
        if rtl:
            set_run_rtl(run)


def clean_markdown(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text).strip()


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    total_twips = 9860
    widths = [total_twips // cols] * cols
    for i, row in enumerate(table.rows):
        if i == 0:
            set_repeat_table_header(row)
        for j, cell in enumerate(row.cells):
            set_cell_width(cell, widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, BRAND_NAVY)
            elif i % 2 == 0:
                set_cell_shading(cell, BRAND_LIGHT)
            p = cell.paragraphs[0]
            set_rtl_paragraph(p)
            add_inline(p, clean_markdown(rows[i][j]))
            for run in p.runs:
                run.font.size = Pt(8.2)
                run.font.color.rgb = RGBColor.from_string(WHITE if i == 0 else BRAND_DARK)
                if i == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF2F3")
    p_pr.append(shd)
    run = p.add_run("\n".join(lines))
    run.font.name = FONT_CODE
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor.from_string(BRAND_DARK)


def extract_toc(markdown: str, limit: int = 80) -> list[tuple[int, str]]:
    headings = []
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if match:
            title = clean_markdown(match.group(2))
            if len(title) <= 150:
                headings.append((len(match.group(1)), title))
    return headings[:limit]


def add_toc_summary(doc: Document, markdown: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_rtl_paragraph(p)
    add_inline(p, "خريطة المحتوى")
    for level, title in extract_toc(markdown):
        p = doc.add_paragraph()
        set_rtl_paragraph(p)
        p.paragraph_format.right_indent = Inches(0.18 * max(0, level - 1))
        r = p.add_run(("• " if level == 1 else "– ") + title)
        r.font.name = FONT_AR
        r.font.size = Pt(9.2 if level == 1 else 8.6)
        r.font.bold = level == 1
        r.font.color.rgb = RGBColor.from_string(BRAND_NAVY if level == 1 else BRAND_DARK)
        set_run_rtl(r)
    doc.add_page_break()


def add_diagram_gallery(doc: Document, diagram_paths: list[Path], selected: list[int] | None = None) -> None:
    selected = selected or list(range(len(diagram_paths)))
    p = doc.add_paragraph(style="Heading 1")
    set_rtl_paragraph(p)
    add_inline(p, "الرسومات التوضيحية الرئيسية")
    captions = [
        "البنية المنطقية وحدود المحرك المحلي",
        "معالج إدخال مشروع بدأ قبل النظام",
        "حوكمة المراحل الموزونة ومعالم التسليم",
        "خط التحليل المحلي الآمن للكراسة",
        "الشجرة الثابتة للفئات والمجموعات والمتطلبات",
        "دورة مراجعة المرشحين والاعتماد البشري",
        "مقارنة إصدار جديد وتحليل أثره",
    ]
    for idx in selected:
        pimg = doc.add_paragraph()
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pimg.add_run()
        picture = run.add_picture(str(diagram_paths[idx]), width=Inches(6.55))
        picture._inline.docPr.set("descr", captions[idx])
        picture._inline.docPr.set("title", f"شكل {idx + 1}")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(f"شكل {idx + 1}: {captions[idx]}")
        rr.font.name = FONT_AR
        rr.font.size = Pt(8.5)
        rr.font.italic = True
        rr.font.color.rgb = RGBColor.from_string(BRAND_MUTED)
        set_run_rtl(rr)
        if idx != selected[-1]:
            doc.add_page_break()
    doc.add_page_break()


def markdown_to_docx(markdown: str, out_path: Path, title: str, subtitle: str, short_title: str, diagrams: list[Path], diagram_selection: list[int]) -> None:
    doc = Document()
    configure_styles(doc)
    set_page_layout(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "CloudTech / Project Desk"
    doc.core_properties.keywords = "Project Desk, Development v2, SRS, Documentation"
    add_cover(doc, title, subtitle)
    add_toc_summary(doc, markdown)
    add_diagram_gallery(doc, diagrams, diagram_selection)

    lines = markdown.splitlines()
    i = 0
    first_h1_seen = False
    in_code = False
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        paragraph_buffer = []
        if text:
            p = doc.add_paragraph()
            set_rtl_paragraph(p)
            add_inline(p, clean_markdown(text))

    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            if level == 1 and first_h1_seen:
                doc.add_page_break()
            first_h1_seen = first_h1_seen or level == 1
            p = doc.add_paragraph(style=f"Heading {level}")
            set_rtl_paragraph(p)
            add_inline(p, clean_markdown(heading.group(2)))
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
                    table_rows.append(cells)
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            set_repeat_table_header(table.rows[0])
            cell = table.cell(0, 0)
            set_cell_width(cell, 9600)
            set_cell_shading(cell, "E8F8F8")
            p = cell.paragraphs[0]
            set_rtl_paragraph(p)
            add_inline(p, " ".join(quote_lines))
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", raw)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", raw)
        checkbox = re.match(r"^\s*[-*]\s+\[([ xX])\]\s+(.+)$", raw)
        if checkbox or bullet or numbered:
            flush_paragraph()
            if checkbox:
                text = ("☒ " if checkbox.group(1).lower() == "x" else "☐ ") + checkbox.group(2)
                p = doc.add_paragraph(style="List Bullet")
            elif numbered:
                text = numbered.group(1)
                p = doc.add_paragraph(style="List Number")
            else:
                text = bullet.group(1)
                p = doc.add_paragraph(style="List Bullet")
            set_rtl_paragraph(p)
            add_inline(p, clean_markdown(text))
            i += 1
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
            r.font.color.rgb = RGBColor.from_string(BRAND_CYAN)
            i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()
    if code_lines:
        add_code_block(doc, code_lines)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            set_rtl_paragraph(paragraph)
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = FONT_AR
            run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_AR)
    setup_headers_footers(doc, short_title)
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_documents(source_paths: dict[str, Path], diagrams: list[Path]) -> list[Path]:
    specs = [
        ("Project-Desk-Development-v2-SRS-IEEE830-AR.md", "Project-Desk-Development-v2-SRS-IEEE830-AR.docx", "كراسة مواصفات متطلبات النظام الشاملة", "SRS وفق IEEE 830 — خط الأساس الكامل مع متطلبات نسخة التطوير", "SRS IEEE 830", [0, 1, 2, 3, 4, 5, 6]),
        ("Project-Desk-Development-v2-Technical-Documentation-AR.md", "Project-Desk-Development-v2-Technical-Documentation-AR.docx", "التوثيق التقني الشامل", "البنية والبيانات والواجهات والأمان والتشغيل والتطوير", "التوثيق التقني", [0, 3, 4, 5]),
        ("Project-Desk-Development-v2-User-Manual-AR.md", "Project-Desk-Development-v2-User-Manual-AR.docx", "دليل المستخدم الشامل", "العمل النموذجي في النسخة الأولى ووظائف نسخة التطوير", "دليل المستخدم", [1, 2, 3, 5, 6]),
        ("Project-Desk-Development-v2-Workflows-AR.md", "Project-Desk-Development-v2-Workflows-AR.docx", "دليل مسارات العمل", "العمليات النموذجية والقرارات والأدوار ورسومات التدفق", "مسارات العمل", [1, 2, 3, 4, 5, 6]),
        ("Project-Desk-Development-v2-Competitive-Analysis-AR.md", "Project-Desk-Development-v2-Competitive-Analysis-AR.docx", "المقارنة التنافسية وتحليل الفجوات", "الدراسة المرجعية الكاملة وأثر إضافات نسخة التطوير", "المقارنة التنافسية", [0, 2, 3, 4]),
    ]
    outputs = []
    for source_name, docx_name, title, subtitle, short_title, selected in specs:
        out_path = OUT / docx_name
        markdown_to_docx(read_text(source_paths[source_name]), out_path, title, subtitle, short_title, diagrams, selected)
        outputs.append(out_path)
        print(f"created {out_path.name}")
    return outputs


def write_manifest() -> Path:
    manifest = OUT / "MANIFEST-SHA256.txt"
    entries = []
    for path in sorted(OUT.rglob("*")):
        if not path.is_file() or path == manifest or path.suffix.lower() == ".zip":
            continue
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        entries.append(f"{digest}  {path.relative_to(OUT).as_posix()}")
    manifest.write_text("\n".join(entries) + "\n", encoding="utf-8")
    return manifest


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    sources = compose_sources()
    diagrams = create_diagrams()
    build_documents(sources, diagrams)
    write_manifest()
    print(f"Package staged at: {OUT}")


if __name__ == "__main__":
    main()
